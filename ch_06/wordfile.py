from pathlib import Path
from datetime import datetime as dt
from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Pt
from docx.text.run import Run
from folder import out_dir

def apply_font_sty(run:Run,size_pt:int=None,is_bold:bool=None):
    if size_pt is not None:
        run.font.size=Pt(size_pt)
    if is_bold is not None:
        run.font.bold=is_bold

def init_docx()->DocumentObject:
    doc=Document()
    p1=doc.add_paragraph(style="Heading1")
    run1=p1.add_run("쇼핑 트렌드 보고서")
    apply_font_sty(run1,size_pt=25,is_bold=True)
    date_string=dt.now().strftime("(%Y.%m.%d)")
    apply_font_sty(p1.add_run(date_string),size_pt=15)
    return doc

if __name__=="__main__":
    doc=init_docx()
    doc.save(out_dir/f"{Path(__file__).stem}.docx")